from docx import Document
from docx.shared import Pt
import os

# 创建 wordFiles 文件夹（如果不存在）
output_dir = "wordFiles"
os.makedirs(output_dir, exist_ok=True)

# 创建文档对象
doc = Document()

# 设置默认字体大小
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# 添加标题
doc.add_heading("完善版 主体-AAA.com(新模板）", level=0)

# 添加基础字段
doc.add_paragraph("URL: \n")
doc.add_paragraph("Payment method：Visa,Mastercard\n")
doc.add_paragraph("Product Type：\n")
doc.add_paragraph("MCC code:\n")
doc.add_paragraph("\n")

# 定义每个审查项结构
sections = [
    ("Terms & Conditions", [
        "REQ:Separate page to be dedicated for T&C",
        "Clearly disclose the conditions of any promotion, discount, or trial that you offer to customers. Display a link or disclaimer text so that it’s visible when customers agree to participate. Transparency around these conditions can help avoid confusion and disputes.",
        "☑Match", "☐No Match",
        "Reason:",
        "☐Missing T&C link",
        "☐Missing content",
        "☐Missing / wrong company details (Company name, Company address and Registration Number)",
        "☐Missing legal information applicable to arbitration",
        "☐The laws and regulations followed are inconsistent with the country where the main company is located",
        "☐Other domain names appear"
    ]),
    ("Privacy Policy", [
        "REQ:Separate page to be dedicated for privacy policy",
        "Consumer data privacy is now a priority for legislation and governments around the world. Clearly explaining your website’s privacy policy helps you both comply with privacy laws and helps your customers understand how their data is protected, used, or disclosed.",
        "☑Match", "☐No Match",
        "Reason:",
        "☐Missing Privacy Policy link",
        "☐Missing content",
        "☐The laws and regulations followed are inconsistent with the country where the main company is located",
        "☐Missing / wrong personal information policy or data protection policy (EU/UK GDPR) of the country where the company is located"
    ]),
    ("Contact Us/Customer Support", [
        "Make sure your customers can find multiple contact methods on your site, including direct communication channels, such as email addresses, phone numbers, and live chat (something besides contact forms). Low-friction communication is key to providing a good customer experience and heading off misunderstandings early on, helping to avoid disputes.",
        "If we review your website and can’t find a clear way to contact you, we may ask that you add some contact options to the site",
        "☑Match", "☐No Match",
        "Reason:",
        "☐Missing Contact Us link",
        "☐Missing content",
        "☐Missing / wrong company details (Company name, Company address and Registration Number)",
        "☐Missing / wrong contact details (Business Email Address or Telephone Number)",
        "☐Wrong domain name in the Business Email Address",
        "☐The laws and regulations followed are inconsistent with the country where the main company is located"
    ]),
    ("Fulfillment Policy (Refund/Cancellation/Return)", [
        "Explanation on your cancellation and refund and return policy for the end-user",
        "Refund policy – Describe the conditions under which customers can receive a refund.",
        "Return policy – Describe the conditions under which customers can return purchased goods.",
        "Cancellation policy – Describe the conditions under which customers can cancel subscriptions or reservations.",
        "☑Match", "☐No Match",
        "Reason:",
        "☐Missing Policy link",
        "☐Missing content",
        "☐Missing / wrong return and exchange conditions",
        "☐Missing / wrong refund conditions",
        "☐Missing / wrong partial refund information"
    ]),
    ("Shipping Policy", [
        "Explanation on your shipment terms and export restrictions",
        "Shipping policy – Describe how and where goods are shipped, and on what timeline.",
        "☑Match", "☐No Match",
        "Reason:",
        "☐Missing Shipping Policy link",
        "☐Missing / wrong content",
        "☐Missing / wrong description of shipping method",
        "☐Missing / wrong description of shipment restriction",
        "☐Missing / wrong description of freight charge standards",
        "☐Missing / wrong description of delivery time",
        "☐Missing / wrong description of free shipping policy"
    ]),
    ("Product and Service Description for Sales", [
        "Product Price, currency, and membership package must be clearly described",
        "☑Match", "☐No Match",
        "Reason:",
        "☐No products",
        "☐IP Rights Infringement",
        "☐Chinese  appears on the product image",
        "☐Selling prohibited goods/services",
        "☐Missing ingredients list"
    ]),
    ("Accepted Payment Methods Logo", [
        "The logos of the credit cards or local payment logo you accept",
        "You can reduce friction in the checkout process by displaying the brand logos of the credit cards that you accept, making it clear to customers that you accept their preferred card.",
        "☑Match", "☐No Match"
    ]),
    ("Website domain name", [
        "☑Match", "☐No Match",
        "Reason:",
        "☐Second-layer domain name",
        "☐Contain non-European countries",
        "☐Restricted words"
    ]),
    ("Additional information", [
        "Additional information is generated when you operate a special business or in a designated area.",
        "☑Match", "☐No Match"
    ]),
    ("Result", [
        "☑Pass", "☐Need further update"
    ]),
    ("Comments", [
        "The website format is correct, there are no illegal products, and it is correct after inspection."
    ]),
    ("DATE", [
        "20250101"
    ])
]

# 写入各段落
for title, content in sections:
    doc.add_paragraph(title, style='Heading 2')
    for line in content:
        doc.add_paragraph(line)

# 保存文档
output_path = os.path.join(output_dir, "recreated_AAA_template.docx")
doc.save(output_path)

print(f"✅ 文档已生成：{output_path}")