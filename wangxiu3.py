
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os
import time
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import NoSuchElementException, ElementClickInterceptedException
import time

# 广告页面弹窗处理
def suspondWindowHandler(browser):
    # 第一种广告弹窗
    try:
        suspondWindow = browser.find_element_by_xpath("//div[contains(@class, 'identity-dialog')]//*[contains(@class, 'close-icon')]")
        suspondWindow.click()
        print(f"searchKey: Suspond Page1 had been closed.")
    except Exception as e:
        print(f"searchKey: there is no suspond Page1. e = {e}")
    # 第二种广告弹窗
    # 如果有广告界面弹出，关闭广告。 否则会导致数据无法输入到搜索框
    try:
        suspondWindow = browser.find_element_by_xpath("//div[contains(@class,'overlay-box')]//div[contains(@class,'overlay-close')]")
        suspondWindow.click()
        print(f"searchKey: Suspond Page2 had been closed.")
    except Exception as e:
        print(f"searchKey: there is no suspond Page2. e = {e}")

def add_picture_if_exists(doc, image_path):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width = Inches(5))
    else:
        doc.add_paragraph('No such file or directory')

def sanitize_filename(name):
    # 去除前后空格
    name = name.strip()
    # 将非法字符替换成下划线
    return re.sub(r'[\\/*?:"<>|]', "_", name).replace(".", "_")

def detect_payment_logos_from_footer_selenium(driver):
    try:
        footer = driver.find_element(By.TAG_NAME, 'footer')
        img_elements = footer.find_elements(By.TAG_NAME, 'img')

        found = []
        known_keywords = ['VISA','visa','Visa',
                          'MASTERCARD','mastercard','Mastercard','MASTER CARD','master card','Master card',
                          'PAYPAL', 'paypal', 'paypal',
                          'APPLEPAY','Applepay', 'applepay','APPLE PAY','Apple pay', 'apple pay',
                          'GOOGLEPAY', 'Googlepay', 'googlepay', 'GOOGLE PAY', 'Google pay', 'google pay',
                          'AMEX','Amex','amex',
                          'AMERICANEXPRESS','americanexpress','AmericanExpress',
                          'KLARNA','klarna','Klarna',
                          'MAESTRO','Maestro','maestro'
                          'PAY','pay','Pay'
                          ]

        for img in img_elements:
            src = img.get_attribute('src') or ""
            alt = img.get_attribute('alt') or ""
            title = img.get_attribute('title') or ""
            combined = (src + alt + title).lower()

            for keyword in known_keywords:
                if keyword in combined:
                    found.append(keyword.capitalize())

        return list(set(found))  # 去重返回
    except Exception as e:
        print(f"⚠️ Failed to detect logos in footer: {e}")
        return []

mcc_lookup = {
    5712: "Furniture, home furnishings and equipment shops and manufacturers, except appliances",
    5719: "Miscellaneous home furnishing speciality shops",
    5722: "Household appliance shops",
    5732: "Electronics shops",
    5733: "Music Stores—Musical Instruments, Pianos, Sheet Music",
    5734: "Computer software outlets",
    5812: "Eating places and restaurants",
    5815: "Digital Goods-Media, Books, Movies, Music",
    5816: "Digital Goods-Games",
    5817: "Digital Goods-Software Applications (excluding games)",
    5941: "Sporting goods shops",
    5945: "Hobby, toy and game shops",
    5947: "Gift, card, novelty and souvenir shops",
    5950: "Glassware and crystal shops",
    5970: "Artist supply and craft shops",
    5977: "Cosmetic shops",
    5995: "Pet shops, pet food and supplies",
    5999: "Miscellaneous and Specialty Retail Stores",
    4112: "Passenger railways",
    4131: "Bus lines",
    4214: "Motor freight carriers and trucking - local and long distance, moving and storage companies and local delivery",
    4215: "Courier services - air and ground and freight forwarders",
    4225: "Public warehousing and storage - farm products, refrigerated goods and household goods",
    5065: "Electrical parts and equipment",
    5072: "Hardware equipment and supplies",
    5200: "Home supply warehouse outlets",
    5251: "Hardware Stores",
    5311: "Department stores",
    5331: "Variety stores",
    5399: "Miscellaneous general merchandise",
    5499: "Miscellaneous food shops - convenience and speciality retail outlets",
    5533: "Automotive parts and accessories outlets",
    5611: "Men’s and boys’ clothing and accessory shops",
    5621: "Women’s ready-to-wear shops",
    5631: "Women’s accessory and speciality shops",
    5641: "Children’s and infants’ wear shops",
    5651: "Family clothing shops",
    5655: "Sports and riding apparel shops",
    5661: "Shoe shops",
    5691: "Men’s and women’s clothing shops",
    5698: "Wig and toupee shops",
    5699: "Miscellaneous apparel and accessory shops",
    7278: "用7399代替",
    7298: "Health and beauty spas",
    7375: "Information Retrieval Services",
    7399: "Business services - not elsewhere classified",
    7512: "Automobile rentals",
    8043: "Opticians, optical goods and eyeglasses",
    7999: "Recreation Services–Not Elsewhere Classified",
    4722: "Travel Agencies and Tour Operations",
    8299: "Schools And Educational Services–Not Elsewhere Classified"
}

# 启动 Chrome 浏览器
def she():
    q1 = Options()
    q1.add_argument('--no-sandbox')
    q1.add_argument('--disable-notifications')
    q1.add_experimental_option('detach', True)
    prefs = {
        "translate_whitelists": {
            "zh": "en",
            "ja": "en",
            "ko": "en",
            "fr": "en",
            "de": "en",
            "es": "en",
            "ru": "en",
            "pt": "en",
            "it": "en",
            "ar": "en",
            "nl": "en",
            "tr": "en",
            "vi": "en",
            "pl": "en",
            "th": "en",
            "sv": "en"
        },
        "translate": {
            "enabled": "true"
        }
    }
    q1.add_experimental_option("prefs", prefs)
    return webdriver.Chrome(service=Service('chromedriver.exe'), options=q1)


print("粘贴数据表（直接把前五列粘贴进来，然后按两次回车）：")
raw_data = []
while True:
    line = input()
    if line.strip() == "":
        break
    raw_data.append(line.strip())

dates, domains, mccs, companies = [], [], [], []
lines = []      # 带 https://www 的格式化网址
unformatted = []  # 原始域名列表

for row in raw_data:
    parts = row.split("\t")  # tab 分隔
    if len(parts) >= 5:
        date, _, domain, mcc, company = parts[:5]
        domain = domain.strip()
        dates.append(date)
        domains.append(domain)
        mccs.append(mcc)
        companies.append(company)
        unformatted.append(domain)
        lines.append(f"https://www.{domain.strip()}/")
    else:
        print(f"⚠️ 行格式错误，已跳过：{row}")


browser = she()
# 遍历每个网址
for index, url in enumerate(lines):
    # 截图目录结构：pictures/<domain>/

    pictures_dir = os.path.join("pictures", unformatted[index])
    os.makedirs(pictures_dir, exist_ok=True)


    browser.get(url)
    browser.set_window_position(0, 0)
    browser.maximize_window()


    # Domain 截屏
    filename = os.path.join(pictures_dir, f"{unformatted[index]}_DOMAIN.png")
    browser.execute_script("document.body.style.zoom='50%'")
    time.sleep(2)
    browser.get_screenshot_as_file(filename)

    # 滑到最后截 payment
    browser.execute_script("window.scrollTo(0,document.body.scrollHeight);")
    filename = os.path.join(pictures_dir, f"{unformatted[index]}_PAYMENT.png")
    time.sleep(2)
    browser.get_screenshot_as_file(filename)

    # Product 页面截两张
    product = ['ALL', 'all', 'All', 'NEW', 'new', 'New', 'SALE', 'sale', 'Sale',
               'SHOP', 'shop', 'Shop','SALE','sale','Sale','MORE','more','More','DRESS','dress','Dress','ARRIVAL',
               'arrival','Arrival','SELL','sell','Sell','COLLECTION','collection','Collection','PRODUCT', 'product', 'Product']
    for text in product:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_PRODUCT1.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(2)
                browser.get_screenshot_as_file(filename)

                browser.execute_script("window.scrollTo(0,document.body.scrollHeight);")
                filename = os.path.join(pictures_dir, f"{unformatted[index]}_PRODUCT2.png")
                time.sleep(2)
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue

    # Terms 页面截两张
    term = ['TERM', 'term', 'Term', 'AGREEMENT', 'Agreement', 'agreement', 'GENERAL', 'General', 'general','CONDITION','condition','Condition']
    for text in term:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_TERMS1.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(1)
                browser.get_screenshot_as_file(filename)

                browser.execute_script("window.scrollTo(0,document.body.scrollHeight);")
                filename = os.path.join(pictures_dir, f"{unformatted[index]}_TERMS2.png")
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue

    # RETURN 页面
    Return = ['RETURN', 'return', 'Return', 'FULFILLMENT', 'fulfillment', 'Fulfillment', 'REFUND', 'refund', 'Refund',
              'CANCEL', 'cancel', 'Cancel','REFUND','refund','Refund']
    for text in Return:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_RETURN.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(1)
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue

    # CONTACT 页面
    contact = ['CONTACT', 'contact', 'Contact', 'SUPPORT', 'support', 'Support', 'CUSTOMER', 'customer', 'Customer']
    for text in contact:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_CONTACT.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(1)
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue

    # PRIVACY 页面
    privacy = ['PRIVACY', 'privacy', 'Privacy', 'DATA', 'data', 'Data']
    for text in privacy:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_PRIVACY.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(1)
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue

    # SHIPPING 页面
    shipping = ['SHIP', 'ship', 'Ship', 'DELIVER', 'Deliver', 'deliver', 'PACKAGE', 'package', 'Package', 'TRACK',
                'track', 'Track']
    for text in shipping:
        try:
            elem = browser.find_element(By.PARTIAL_LINK_TEXT, text)
            if elem.tag_name.lower() == 'a' and elem.get_attribute('href'):
                browser.execute_script("arguments[0].click();", elem)

                filename = os.path.join(pictures_dir, f"{unformatted[index]}_SHIPPING.png")
                browser.execute_script("document.body.style.zoom='50%'")
                time.sleep(1)
                browser.get_screenshot_as_file(filename)

                browser.get(url)
                break
        except NoSuchElementException:
            continue


    payment_logos = detect_payment_logos_from_footer_selenium(browser)
    print(f"detected payment logos:{payment_logos}")
    # ---------------------------------------------------------------------
    # 创建 wordFiles 文件夹（如果不存在）
    output_dir = "wordFiles"
    os.makedirs(output_dir, exist_ok=True)

    # 创建文档对象
    doc = Document()

    # 设置默认字体大小
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # 添加标题
    doc.add_heading(f"{companies[index]}-{unformatted[index]}", level=0)

    # 添加基础字段
    doc.add_paragraph(f"URL: {unformatted[index]}\n")
    doc.add_paragraph("Payment method：Visa,Mastercard\n")
    doc.add_paragraph(f"Product Type：{mcc_lookup.get(int(mccs[index]),'unknown')}\n")
    doc.add_paragraph(f"MCC code: {mccs[index]}\n")
    doc.add_paragraph("\n")

    # 定义每个审查项结构
    from docx import Document
    from docx.shared import Pt, Inches
    import os

    # 创建文档对象
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # 添加标题
    doc.add_heading(f"{companies[index]}-{unformatted[index]}", level=0)

    # 添加基础字段
    doc.add_paragraph(f"URL: {unformatted[index]}\n")
    doc.add_paragraph("Payment method：Visa,Mastercard\n")
    doc.add_paragraph(f"Product Type：{mcc_lookup.get(int(mccs[index]), 'unknown')}\n")
    doc.add_paragraph(f"MCC code: {mccs[index]}\n")
    doc.add_paragraph("\n")


    if payment_logos:
        payment_match_line = "☑Match"
        payment_nomatch_line =  "☐No Match"
    else:
        payment_match_line = "☑Match"
        payment_nomatch_line = "☐No Match"
    # 模拟插图段落结构
    sections = [
        ("Terms & Conditions", [
            "REQ:Separate page to be dedicated for T&C",
            "Clearly disclose the conditions of any promotion, discount, or trial that you offer to customers. Display a link or disclaimer text so that it’s visible when customers agree to participate. Transparency around these conditions can help avoid confusion and disputes.",
            "☑Match", "☐No Match",
            "Reason:",
            "☐Missing T&C link",
            "☐Missing content",
            "☐Missing / wrong company details (Company name, Company address and Registration Number)",
            "☐Missing legal information applicable to arbitration",
            "☐The laws and regulations followed are inconsistent with the country where the main company is located",
            "☐Other domain names appear"
        ], [f"{unformatted[index]}_TERMS1.png", f"{unformatted[index]}_TERMS2.png"]),
        ("Privacy Policy", [
            "REQ:Separate page to be dedicated for privacy policy",
            "Consumer data privacy is now a priority for legislation and governments around the world. Clearly explaining your website’s privacy policy helps you both comply with privacy laws and helps your customers understand how their data is protected, used, or disclosed.",
            "☑Match", "☐No Match",
            "Reason:",
            "☐Missing Privacy Policy link",
            "☐Missing content",
            "☐The laws and regulations followed are inconsistent with the country where the main company is located",
            "☐Missing / wrong personal information policy or data protection policy (EU/UK GDPR) of the country where the company is located"
        ], [f"{unformatted[index]}_PRIVACY.png"]),
        ("Contact Us/Customer Support", [
            "Make sure your customers can find multiple contact methods on your site, including direct communication channels, such as email addresses, phone numbers, and live chat (something besides contact forms). Low-friction communication is key to providing a good customer experience and heading off misunderstandings early on, helping to avoid disputes.",
            "If we review your website and can’t find a clear way to contact you, we may ask that you add some contact options to the site",
            "☑Match", "☐No Match",
            "Reason:",
            "☐Missing Contact Us link",
            "☐Missing content",
            "☐Missing / wrong company details (Company name, Company address and Registration Number)",
            "☐Missing / wrong contact details (Business Email Address or Telephone Number)",
            "☐Wrong domain name in the Business Email Address",
            "☐The laws and regulations followed are inconsistent with the country where the main company is located"
        ], [f"{unformatted[index]}_CONTACT.png"]),
        ("Fulfillment Policy (Refund/Cancellation/Return)", [
            "Explanation on your cancellation and refund and return policy for the end-user",
            "Refund policy – Describe the conditions under which customers can receive a refund.",
            "Return policy – Describe the conditions under which customers can return purchased goods.",
            "Cancellation policy – Describe the conditions under which customers can cancel subscriptions or reservations.",
            "☑Match", "☐No Match",
            "Reason:",
            "☐Missing Policy link",
            "☐Missing content",
            "☐Missing / wrong return and exchange conditions",
            "☐Missing / wrong refund conditions",
            "☐Missing / wrong partial refund information"
        ], [f"{unformatted[index]}_RETURN.png"]),
        ("Shipping Policy", [
            "Explanation on your shipment terms and export restrictions",
            "Shipping policy – Describe how and where goods are shipped, and on what timeline.",
            "☑Match", "☐No Match",
            "Reason:",
            "☐Missing Shipping Policy link",
            "☐Missing / wrong content",
            "☐Missing / wrong description of shipping method",
            "☐Missing / wrong description of shipment restriction",
            "☐Missing / wrong description of freight charge standards",
            "☐Missing / wrong description of delivery time",
            "☐Missing / wrong description of free shipping policy"
        ], [f"{unformatted[index]}_SHIPPING.png"]),
        ("Product and Service Description for Sales", [
            "Product Price, currency, and membership package must be clearly described",
            "☑Match", "☐No Match",
            "Reason:",
            "☐No products",
            "☐IP Rights Infringement",
            "☐Chinese  appears on the product image",
            "☐Selling prohibited goods/services",
            "☐Missing ingredients list"
        ], [f"{unformatted[index]}_PRODUCT1.png", f"{unformatted[index]}_PRODUCT2.png"]),
        ("Accepted Payment Methods Logo", [
            "The logos of the credit cards or local payment logo you accept",
            "You can reduce friction in the checkout process by displaying the brand logos of the credit cards that you accept, making it clear to customers that you accept their preferred card.",
            payment_match_line,
            payment_nomatch_line
        ], [f"{unformatted[index]}_PAYMENT.png"]),
        ("Website domain name", [
            "☑Match", "☐No Match",
            "Reason:",
            "☐Second-layer domain name",
            "☐Contain non-European countries",
            "☐Restricted words"
        ], [f"{unformatted[index]}_DOMAIN.png"]),
        ("Additional information", [
            "Additional information is generated when you operate a special business or in a designated area.",
            "☑Match", "☐No Match"
        ]),
        ("Result", [
            "☑Pass", "☐Need further update"
        ]),
        ("Comments", [
            "The website format is correct, there are no illegal products, and it is correct after inspection."
        ]),
        ("DATE", [
            f"{dates[index]}"
        ])
    ]

    # 插入内容和图片
    for section in sections:
        # 解包每个 section 的结构
        if len(section) == 3:
            title, content_lines, image_files = section
        else:
            title, content_lines = section
            image_files = []

        # 添加标题
        doc.add_paragraph(title, style='Heading 2')

        # 添加每行内容
        for line in content_lines:
            doc.add_paragraph(line)

        # 添加图片（如有）
        for img_file in image_files:
            img_path = os.path.join(pictures_dir, img_file)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5))
            else:
                doc.add_paragraph(f"(⚠️ Image not found: {img_file})")




    # 保存文档
    output_path = os.path.join(output_dir, f"{companies[index]}-{unformatted[index]}.docx")
    doc.save(output_path)

    print(f"✅ 文档已生成：{output_path}")





    #---------------------------------------------------------------------
browser.quit()
print("记得检查一遍再用wordToPdf转译")
